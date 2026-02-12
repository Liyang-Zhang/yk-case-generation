"""Parse DOCX attachments into Source with line-level evidence.

Notes:
- Avoid template assumptions; keep raw text order (paragraphs, then tables in document order).
- Checkbox-like glyphs (e.g., □, √, ✓) are preserved in text; flagged for downstream awareness.
"""
from pathlib import Path
from typing import List
from docx import Document

from yk_case_generation.models.document_ir import Source, Page, Line

_CHECKBOX_CHARS = {"□", "■", "☑", "☐", "√", "✓", "✗", "✘"}


def parse_docx(path: Path) -> Source:
    """Parse DOCX to a text-only Source (no layout), minimizing noise."""
    doc = Document(path)
    lines: List[Line] = []
    line_id = 1

    # Paragraphs in order
    for idx, para in enumerate(doc.paragraphs, start=1):
        text = para.text
        if not text.strip() and not _has_checkbox_char(text):
            continue
        flags = {}
        if _has_checkbox_char(text):
            flags["checkbox_like"] = True
        lines.append(
            Line(
                line_id=line_id,
                text=text,
                confidence=None,
                parag_no=idx,
                flags=flags,
            )
        )
        line_id += 1

    # Tables (row-major order). Skip empty cells and collapse perfect duplicates in the same row.
    for t_idx, table in enumerate(doc.tables, start=1):
        for r_idx, row in enumerate(table.rows):
            last_text = None
            for c_idx, cell in enumerate(row.cells):
                text = cell.text
                if not text.strip() and not _has_checkbox_char(text):
                    continue
                if last_text is not None and text == last_text:
                    # suppress exact duplicates caused by merged cells repeats
                    last_text = text
                    continue
                flags = {"table": True, "table_index": t_idx, "row": r_idx, "col": c_idx}
                if _has_checkbox_char(text):
                    flags["checkbox_like"] = True
                lines.append(
                    Line(
                        line_id=line_id,
                        text=text,
                        confidence=None,
                        parag_no=None,
                        flags=flags,
                    )
                )
                line_id += 1
                last_text = text

    # Inline images count (hint for missed checkboxes)
    inline_shapes = getattr(doc, "inline_shapes", [])
    if inline_shapes:
        lines.append(
            Line(
                line_id=line_id,
                text=f"[inline_shapes:{len(inline_shapes)}]",
                confidence=None,
                parag_no=None,
                flags={"inline_shapes_count": len(inline_shapes)},
            )
        )

    page = Page(page_number=None, lines=lines)
    return Source(source_id=path.stem + "_docx_text", source_type="docx", pages=[page])


def _has_checkbox_char(text: str) -> bool:
    return any(ch in _CHECKBOX_CHARS for ch in text)
